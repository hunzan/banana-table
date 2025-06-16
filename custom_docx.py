from docx import Document
from io import BytesIO
from docx.shared import Pt, Mm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT

def generate_custom_docx(titles, contents, table_title="自訂表格", border_width="8", font_name="DFKai-SB", layout="portrait"):
    doc = Document()

    # 設定頁面方向
    section = doc.sections[0]  # 取得第一個 section
    if layout == "landscape":
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Mm(297)   # A4 寬度
        section.page_height = Mm(210)  # A4 高度
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        
    # 標題
    heading = doc.add_heading(level=0)
    run = heading.add_run(table_title)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')  # 英文 fallback
    run.font.size = Pt(20)
    run.bold = True
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 表格
    table = doc.add_table(rows=0, cols=1)
    set_table_borders(table, border_width)  # 👈 傳入粗細參數
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    for title, content in zip(titles, contents):
        # 項目列
        row_label = table.add_row().cells[0]
        p_label = row_label.paragraphs[0]
        run_label = p_label.add_run(title)
        run_label.font.name = font_name
        run_label._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run_label._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')  # 英文 fallback
        run_label.font.size = Pt(12)
        run_label.bold = True
        p_label.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 內容列
        row_content = table.add_row().cells[0]
        p_content = row_content.paragraphs[0]

        for line in content.splitlines():
            run_line = p_content.add_run(line)
            run_line.font.name = font_name
            run_line._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            run_line._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')  # 英文 fallback
            run_line.font.size = Pt(12)
            p_content.add_run().add_break()  # 每行後換行


    # ✅ 把 doc 存進 buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def set_table_borders(table, border_width="8"):
    tbl = table._tbl
    tblPr = tbl.tblPr

    borders = OxmlElement('w:tblBorders')

    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), border_width)  # 👈 自訂粗細
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        borders.append(border)

    tblPr.append(borders)