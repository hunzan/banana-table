import io
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.enum.table import WD_ALIGN_VERTICAL
import docx.oxml

def generate_schedule_docx(data):
    font_choice = data.get('font', 'DFKai-SB')  # 預設標楷體
    orientation = data.get('orientation', 'landscape')
    border_width = str(data.get("borderWidth", 8))  # 預設 8（1pt）
    doc = Document()

    # 設定橫式或直式 A4
    section = doc.sections[0]
    if orientation == 'landscape':
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = Inches(11.69), Inches(8.27)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width, section.page_height = Inches(8.27), Inches(11.69)

    # 標題（加大字體、置中、加粗）
    if data.get("title"):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(data["title"])
        run.font.name = font_choice
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_choice)
        run.font.size = Pt(20)
        run.bold = True

    # 說明（表格前）
    if data.get("remarks"):
        para = doc.add_paragraph()
        run = para.add_run(data["remarks"])
        run.font.name = font_choice
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_choice)

    weeks = data.get("weeks", [])
    periods = data.get("periods", [])
    times = data.get("times", [])
    content = data.get("content", [])
    breaks = data.get("breaks", [])
    note = data.get("note", "")

    break_dict = {}
    for b in breaks:
        try:
            idx, text = b.split(":")
            break_dict[int(idx.strip())] = text.strip()
        except ValueError:
            continue

    # 表頭 + 空資料列，後面補進內容
    table = doc.add_table(rows=1, cols=len(weeks) + 1)
    table.style = 'Table Grid'

    # 表頭
    hdr_cells = table.rows[0].cells
    hdr_para = hdr_cells[0].paragraphs[0]
    hdr_run = hdr_para.add_run("節次／時間")
    hdr_run.bold = True
    hdr_run.font.name = font_choice
    hdr_run._element.rPr.rFonts.set(qn('w:eastAsia'), font_choice)

    for i, week in enumerate(weeks):
        para = hdr_cells[i + 1].paragraphs[0]
        run = para.add_run(week)
        run.bold = True
        run.font.name = font_choice
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_choice)

    # 建立 row map: 課程列與 break 列插入邏輯
    row_map = []
    for i in range(len(periods)):
        row_map.append(('period', i))
        if (i + 1) in break_dict:
            row_map.append(('break', i + 1))

    # 建立所有資料列
    for _ in range(len(row_map)):
        table.add_row()

    # 寫入內容
    for row_idx, (row_type, val) in enumerate(row_map):
        cells = table.rows[row_idx + 1].cells  # +1 因為表頭是第0列

        if row_type == 'period':
            time_str = times[val] if val < len(times) else ""
            para = cells[0].paragraphs[0]
            run = para.add_run(f"{periods[val]}\n{time_str}")
            run.font.name = font_choice
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_choice)

            for j in range(len(weeks)):
                try:
                    cell_text = content[val][j]
                except IndexError:
                    cell_text = ""
                para = cells[j + 1].paragraphs[0]
                run = para.add_run(cell_text)
                run.font.name = font_choice
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_choice)

        elif row_type == 'break':
            merged_cell = cells[0]
            for k in range(1, len(cells)):
                merged_cell = merged_cell.merge(cells[k])

            merged_cell.text = break_dict[val]
            merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for para in merged_cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.name = font_choice
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_choice)

    # 備註
    if note:
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(12)
        run = para.add_run(f"備註：{note}")
        run.font.name = font_choice
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_choice)

    # 表格樣式統一：文字置中
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.name = font_choice
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_choice)

            # 👉 加這段設定框線粗細
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            for border_tag in ['top', 'left', 'bottom', 'right']:
                element = tcPr.find(qn(f'w:{border_tag}'))
                if element is None:
                    element = docx.oxml.OxmlElement(f'w:{border_tag}')
                    tcPr.append(element)
                element.set(qn('w:val'), 'single')
                element.set(qn('w:sz'), border_width)
                element.set(qn('w:space'), '0')
                element.set(qn('w:color'), 'auto')

    # 輸出
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io
