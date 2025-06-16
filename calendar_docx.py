# calendar_docx.py
import io
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import calendar
from docx.oxml.ns import qn
from docx.enum.section import WD_ORIENT
from datetime import datetime

def generate_monthly_calendar_docx(data):
    title = data.get("title", "月行事曆")
    year = int(data.get("year"))
    month = int(data.get("month"))
    layout = data.get("layout", "portrait")  # A4直式或橫式
    font = data.get("font", "標楷體")
    border_width = data.get("borderWidth", "8")
    fixed_events = data.get("fixedEvents", [])
    specific_events = data.get("specificEvents", [])

    # 🟢 只在日期不是 YYYY-MM-DD 時，才嘗試轉換
    for event in specific_events:
        try:
            old_date = event["date"]
            # 若已經是 YYYY-MM-DD 就跳過
            if len(old_date) == 10 and old_date.count("-") == 2:
                continue
            # 嘗試從 MM/DD 轉成 datetime
            date_obj = datetime.strptime(old_date, "%m/%d")
            date_obj = date_obj.replace(year=year)
            event["date"] = date_obj.strftime("%Y-%m-%d")
        except Exception as e:
            print(f"⚠️ 特定行程日期格式錯誤：{event} → {e}")
    
    # 若 fixed_events 是字串陣列，轉成 dict 陣列
    if fixed_events and isinstance(fixed_events[0], str):
        new_fixed_events = []
        for item in fixed_events:
            try:
                # 假設格式是「每週一 下午3點 居服來訪」
                parts = item.replace("每週", "").split(" ")
                new_fixed_events.append({
                    "weekday": parts[0],
                    "time": parts[1],
                    "task": " ".join(parts[2:])  # 有可能有多個詞
                })
            except Exception as e:
                print(f"⚠️ 固定行程格式錯誤：{item} → {e}")
        fixed_events = new_fixed_events

    doc = Document()
    # ✅ 套用整份文件的預設字型
    style = doc.styles['Normal']
    style.font.name = font
    style.font.size = Pt(10)

    # 取第一個 Section 來設定紙張方向與大小
    section = doc.sections[0]

    if layout == "landscape":
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
    else:
        section.orientation = WD_ORIENT.PORTRAIT

    heading = doc.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER  # ✅ 顯式置中
    # ✅ 額外指定標題字型大小
    for run in heading.runs:
        run.font.name = font
        run.font.size = Pt(20)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font)
        heading.runs[0].bold = True

    # 月曆表格（以週為列，日為欄）
    cal = calendar.Calendar()
    month_days = cal.monthdayscalendar(year, month)

    table = doc.add_table(rows=len(month_days)+1, cols=7)
    table.style = 'Table Grid'

    # 設定表格框線粗細
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()

            # 建立四邊框線設定
            for border_dir in ["top", "bottom", "left", "right"]:
                border = tcPr.find(qn(f'w:{border_dir}'))
                if border is None:
                    border = tcPr.makeelement(qn(f'w:{border_dir}'))
                    tcPr.append(border)
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), str(int(border_width) * 8))  # Word 內部單位是 1/8 pt
                border.set(qn('w:space'), "0")
                border.set(qn('w:color'), "000000")


    # 表頭（星期幾）
    headers = ['一', '二', '三', '四', '五', '六', '日']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.height = Inches(1.0)
        run = para.runs[0]
        run.font.name = font
        run.font.size = Pt(12)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font)
        run.font.bold = True

    # 各日格子
        # 先建立一個週幾對照表
    weekday_map = {
        0: "一", 1: "二", 2: "三", 3: "四", 4: "五", 5: "六", 6: "日"
    }

    for row_idx, week in enumerate(month_days):
        for col_idx, day in enumerate(week):
            cell = table.cell(row_idx+1, col_idx)
            content = ""

            if day != 0:
                content = f"{day}"

                full_date = f"{year}-{month:02}-{day:02}"

                # 加入非固定行程
                daily_events = [e for e in specific_events if e.get("date") == full_date]
                for event in daily_events:
                    content += f"\n{event.get('time', '')} {event.get('task', '')}"

                # 加入固定行程
                weekday_ch = weekday_map[col_idx]
                daily_fixed = [e for e in fixed_events if e.get("weekday") == weekday_ch]
                for fevent in daily_fixed:
                    content += f"\n{fevent.get('time', '')} {fevent.get('task', '')}"

            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

            if day != 0:
                # 日期數字（粗體）
                run = para.add_run(f"{day}\n")
                run.font.name = font
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font)
                run.font.size = Pt(12)
                run.font.bold = True

                # 非固定行程
                for event in daily_events:
                    run = para.add_run(f"{event.get('time', '')} {event.get('task', '')}\n")
                    run.font.name = font
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), font)
                    run.font.size = Pt(9)

                # 固定行程
                for fevent in daily_fixed:
                    run = para.add_run(f"{fevent.get('time', '')} {fevent.get('task', '')}\n")
                    run.font.name = font
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), font)
                    run.font.size = Pt(9)
            else:
                # 空白格子也塞進三個換行（含全形空格），避免列高度縮水
                for _ in range(3):
                    run = para.add_run("　\n")  # 全形空格 + 換行
                    run.font.name = font
                    run.font.size = Pt(9)

            # 格子文字輸入後，設定靠左與字體
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT  # ✅ 靠左對齊
            run = para.runs[0]
            run.font.name = font
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font)
            run.font.size = Pt(9)

    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io
